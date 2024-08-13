import tkinter as tk
from tkinter import filedialog, messagebox
import tabula
from docx import Document
import csv
import os

def get_filename_without_extension(file_path):
    # Obtiene el nombre del archivo con la extensión
    file_name_with_extension = os.path.basename(file_path)
    # Separa el nombre del archivo de su extensión
    file_name, _ = os.path.splitext(file_name_with_extension)
    return file_name

def extract_pdf_tables(filepath, output_folder):
    try:
        filename = get_filename_without_extension(filepath)
        print(filename)
        tables = tabula. read_pdf(filepath, pages='all', multiple_tables=True)
        for i, table in enumerate(tables):
            table.to_csv(os.path.join(output_folder, f'{filename}_{i}.csv'), index=False)
    except Exception as err:
        print(f"Unexpected {err=}, {type(err)=}")
        raise

def extract_docx_tables(filepath, output_folder):
    try:
        doc = Document(filepath)
        filename = get_filename_without_extension(filepath)
        for i, table in enumerate(doc.tables):
            with open(os.path.join(output_folder, f'{filename}_{i}.csv'), 'w', newline='') as f:
                writer = csv.writer(f)
                for row in table.rows:
                    writer.writerow([cell.text for cell in row.cells])
    except Exception as err:
        print(f"Unexpected {err=}, {type(err)=}")
        raise

#def extract_pdf_tables(filepath, output_folder):
#
#    print("try to read at least one, with camelot")
#    try: 
#        tables = camelot.read_pdf(filepath, pages='all')
#    except Exception as err:
#        print(f"Unexpected {err=}, {type(err)=}")
#        raise
#
#    for i, table in enumerate(tables):
#        print("try to read at least one")
#        table.to_csv(os.path.join(output_folder, f'tabla_{i}.csv'))

#def extract_docx_tables(filepath, output_folder):
#    doc = Document(filepath)
#    for i, table in enumerate(doc.tables):
#        with open(os.path.join(output_folder, f'tabla_{i}.csv'), 'w', newline='') as f:
#            writer = csv.writer(f)
#            for row in table.rows:
#                writer.writerow([cell.text for cell in row.cells])

def process_files(input_folder, output_folder):
    for filename in os.listdir(input_folder):
        print(f'processing {filename}')
        if os.path.isfile(filename) and (filename.lower().endswith('.pdf') or filename.lower().endswith('.docx')):
            filepath = os.path.join(input_folder, filename)
            try:
                print(f'processing {filepath}')
                if filename.lower().endswith('.pdf'):
                    extract_pdf_tables(filepath, output_folder)
                elif filename.lower().endswith('.docx'):
                    extract_docx_tables(filepath, output_folder)
            except Exception as e:
                print(f"Error processing {filename}: {e}")

def select_folder():
    input_folder = filedialog.askdirectory(title="Select input folder")
    print(f'input folder:{input_folder}')
    if input_folder:
        output_folder = filedialog.askdirectory(title="Select output folder")
        print(f'output folder:{input_folder}')
        if output_folder:
            try:
                process_files(input_folder, output_folder)
                messagebox.showinfo("Success", "Files successfully processed and tables extracted!")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")
        else:
            messagebox.showwarning("No folder selected", "Please select an output folder.")
    else:
        messagebox.showwarning("No folder selected", "Please select an input folder.")

def select_file():
    filetypes = (
        ('PDF files', '*.pdf'),
        ('Word files', '*.docx'),
        ('All files', '*.*')
    )
    filepath = filedialog.askopenfilename(title="Open file", filetypes=filetypes)
    if filepath:
        output_folder = filedialog.askdirectory(title="Select output folder")
        if output_folder:
            try:
                if filepath.lower().endswith('.pdf'):
                    extract_pdf_tables(filepath, output_folder)
                elif filepath.lower().endswith('.docx'):
                    extract_docx_tables(filepath, output_folder)
                messagebox.showinfo("Success", "File successfully processed and tables extracted!")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")
        else:
            messagebox.showwarning("No folder selected", "Please select an output folder.")
    else:
        messagebox.showwarning("No file selected", "Please select a file.")

# Crear la ventana principal
root = tk.Tk()
root.title("Table Extractor")

# Botón para seleccionar la carpeta
select_folder_button = tk.Button(root, text="Select Folder", command=select_folder)
select_folder_button.pack(pady=20)

# Botón para seleccionar el archivo
select_file_button = tk.Button(root, text="Select File", command=select_file)
select_file_button.pack(pady=20)


# Ejecutar la aplicación
root.mainloop()
